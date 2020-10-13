import docx
from seq2seqText import *


def read_docx(docx_file, re_match_file, seq2seq_model):
    file = docx.Document(docx_file)
    re_list = []
    str_list = []
    # 将聚类得到的正则匹配结果和seq2seq语料分别保存
    with open(re_match_file, 'r', encoding='utf-8') as f:
        line = f.readline()
        while line:
            re_temp, str_temp = line.strip().split(',')
            re_list.append(re_temp)
            str_list.append(str_temp)
            line = f.readline()

    translate_str = []
    para_list = []
    for table in file.tables:
        for row in table.rows:
            for i, temp_re in enumerate(re_list):
                if re.search(temp_re, row.cells[2].text) is not None:
                    translate_str.append(seq2seq_model.translate(str_list[i]))
                    break

    # for para in file.paragraphs:
    #     para_list.append(para.text)
    #     for i, temp_re in enumerate(re_list):
    #         if re.search(temp_re, para.text) is not None:
    #             translate_str.append(seq2seq_model.translate(str_list[i]))
    #             break
    return translate_str


if __name__ == '__main__':
    seq2seq_model = Seq2seq()
    translate_result = read_docx('../file/1.docx', '../file/re_match', seq2seq_model)
    print(translate_result)